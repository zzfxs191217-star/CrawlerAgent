"""报告导出：把 Markdown 报告渲染为 PDF / Word 文件。

用法（代码内）：
    from crawler.export import export_report
    export_report(report_md, "报告名", out_dir, fmts=["pdf", "docx"])

用法（命令行）：
    uv run python -m crawler.export --file reports/xx.md --out reports --fmt pdf,docx
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

# PDF 中文字体候选：(正常体, 粗体)，命中第一个可用组合；全部失败则回退内置 STSong-Light
PDF_FONT_CANDIDATES = [
    (r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\msyhbd.ttc"),
    (r"C:\Windows\Fonts\simsun.ttc", r"C:\Windows\Fonts\simhei.ttf"),
    (r"/System/Library/Fonts/PingFang.ttc", None),
]

_BLOCK_RE = re.compile(r"^(#{1,3}\s|>\s?|[-*]\s+|\d+[.)]\s+|\|)")
_INLINE_PART = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\([^)]+\))")


def _starts_block(line: str) -> bool:
    return bool(_BLOCK_RE.match(line.lstrip()))


def _split_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def parse_blocks(md: str) -> list[dict]:
    """把报告 Markdown 解析为结构块列表：h1/h2/h3/p/quote/ul/ol/table。"""
    lines = md.splitlines()
    blocks: list[dict] = []
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("### "):
            blocks.append({"type": "h3", "text": line[4:].strip()})
            i += 1
        elif line.startswith("## "):
            blocks.append({"type": "h2", "text": line[3:].strip()})
            i += 1
        elif line.startswith("# "):
            blocks.append({"type": "h1", "text": line[2:].strip()})
            i += 1
        elif line.startswith(">"):
            quote = []
            while i < n and lines[i].startswith(">"):
                quote.append(lines[i].lstrip("> ").strip())
                i += 1
            blocks.append({"type": "quote", "text": " ".join(quote)})
        elif line.startswith("|"):
            rows = []
            while i < n and lines[i].startswith("|"):
                rows.append(_split_table_row(lines[i]))
                i += 1
            rows = [
                row
                for row in rows
                if not all(re.fullmatch(r":?-{2,}:?", cell.strip()) for cell in row)
            ]
            if rows:
                blocks.append({"type": "table", "rows": rows})
        elif re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < n and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]).strip())
                i += 1
            blocks.append({"type": "ul", "items": items})
        elif re.match(r"^\s*\d+[.)]\s+", line):
            items = []
            while i < n and re.match(r"^\s*\d+[.)]\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+[.)]\s+", "", lines[i]).strip())
                i += 1
            blocks.append({"type": "ol", "items": items})
        else:
            para = []
            while (
                i < n
                and lines[i].strip()
                and not _starts_block(lines[i])
            ):
                para.append(lines[i].strip())
                i += 1
            blocks.append({"type": "p", "text": " ".join(para)})
    return blocks


def _escape_html(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline_to_html(text: str) -> str:
    """把 **加粗** 与 [文字](链接) 转成 reportlab Paragraph 支持的迷你 HTML。"""
    text = _escape_html(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2" color="#0563C1">\1</a>', text)
    return text


# ---------------- PDF ----------------

def _register_pdf_fonts() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont

    for normal, bold in PDF_FONT_CANDIDATES:
        if not Path(normal).exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("CrawlerCJK", normal))
            bold_name = "CrawlerCJK"
            if bold and Path(bold).exists():
                try:
                    pdfmetrics.registerFont(TTFont("CrawlerCJKBold", bold))
                    bold_name = "CrawlerCJKBold"
                except Exception:
                    pass
            pdfmetrics.registerFontFamily(
                "CrawlerCJK", normal="CrawlerCJK", bold=bold_name,
                italic="CrawlerCJK", boldItalic=bold_name,
            )
            return "CrawlerCJK"
        except Exception:
            continue
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    pdfmetrics.registerFontFamily(
        "CrawlerCJK", normal="STSong-Light", bold="STSong-Light",
        italic="STSong-Light", boldItalic="STSong-Light",
    )
    return "STSong-Light"


def export_pdf(markdown_text: str, out_path: Path) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font = _register_pdf_fonts()
    styles = {
        "h1": ParagraphStyle("h1", fontName=font, fontSize=17, leading=23, spaceAfter=6),
        "h2": ParagraphStyle("h2", fontName=font, fontSize=14, leading=19, spaceBefore=10, spaceAfter=4),
        "h3": ParagraphStyle("h3", fontName=font, fontSize=12, leading=17, spaceBefore=8, spaceAfter=3),
        "p": ParagraphStyle("p", fontName=font, fontSize=10.5, leading=16),
        "quote": ParagraphStyle("quote", fontName=font, fontSize=10, leading=15,
                                leftIndent=10, textColor=colors.HexColor("#555555")),
        "li": ParagraphStyle("li", fontName=font, fontSize=10.5, leading=16,
                             leftIndent=14, bulletIndent=2),
        "cell": ParagraphStyle("cell", fontName=font, fontSize=9.5, leading=13),
    }
    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm,
        title="CrawlerAgent 分析报告", author="CrawlerAgent",
    )
    story = []
    for block in parse_blocks(markdown_text):
        t = block["type"]
        if t in ("h1", "h2", "h3", "p", "quote"):
            story.append(Paragraph(_inline_to_html(block["text"]), styles[t]))
        elif t in ("ul", "ol"):
            for idx, item in enumerate(block["items"], start=1):
                bullet = "•" if t == "ul" else f"{idx}."
                story.append(Paragraph(_inline_to_html(item), styles["li"], bulletText=bullet))
        elif t == "table":
            data = [
                [Paragraph(_inline_to_html(cell), styles["cell"]) for cell in row]
                for row in block["rows"]
            ]
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF1F5")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(table)
        story.append(Spacer(1, 4))
    doc.build(story)
    return out_path


# ---------------- Word ----------------

def _add_hyperlink(paragraph, text: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_rich_text(paragraph, text: str) -> None:
    """把 **加粗** 与 [文字](链接) 写进 Word 段落（普通文字为普通 run）。"""
    for part in _INLINE_PART.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            continue
        m = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", part)
        if m:
            _add_hyperlink(paragraph, m.group(1), m.group(2))
            continue
        paragraph.add_run(part)


def export_docx(markdown_text: str, out_path: Path) -> Path:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    for block in parse_blocks(markdown_text):
        t = block["type"]
        if t == "h1":
            _add_rich_text(doc.add_heading("", level=0), block["text"])
        elif t == "h2":
            _add_rich_text(doc.add_heading("", level=1), block["text"])
        elif t == "h3":
            _add_rich_text(doc.add_heading("", level=2), block["text"])
        elif t == "quote":
            p = doc.add_paragraph()
            _add_rich_text(p, block["text"])
            p.italic = True
        elif t == "p":
            _add_rich_text(doc.add_paragraph(), block["text"])
        elif t == "ul":
            for item in block["items"]:
                _add_rich_text(doc.add_paragraph(style="List Bullet"), item)
        elif t == "ol":
            for item in block["items"]:
                _add_rich_text(doc.add_paragraph(style="List Number"), item)
        elif t == "table":
            rows = block["rows"]
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = table.cell(ri, ci)
                    _add_rich_text(cell.paragraphs[0], cell_text)
            doc.add_paragraph()
    doc.save(str(out_path))
    return out_path


# ---------------- 统一入口 ----------------

def export_report(markdown_text: str, base_name: str, out_dir: Path,
                  fmts: list[str] | None = None) -> dict[str, Path]:
    """把 Markdown 报告导出为指定格式，返回 {格式: 文件路径}。fmts 元素：pdf/docx。"""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    fmts = fmts or ["pdf", "docx"]
    exported: dict[str, Path] = {}
    for fmt in fmts:
        fmt = fmt.strip().lower()
        if fmt == "pdf":
            exported["pdf"] = export_pdf(markdown_text, out_dir / f"{base_name}.pdf")
        elif fmt == "docx":
            exported["docx"] = export_docx(markdown_text, out_dir / f"{base_name}.docx")
        else:
            raise ValueError(f"不支持的导出格式：{fmt}")
    return exported


def main() -> int:
    parser = argparse.ArgumentParser(description="CrawlerAgent 报告导出（Markdown → PDF/Word）")
    parser.add_argument("--file", required=True, help="要导出的 Markdown 报告路径")
    parser.add_argument("--out", default="reports", help="输出目录，默认 reports")
    parser.add_argument("--fmt", default="pdf,docx", help="导出格式，逗号分隔：pdf/docx")
    args = parser.parse_args()
    md_path = Path(args.file)
    if not md_path.exists():
        print(f"文件不存在：{md_path}")
        return 1
    fmts = [f.strip() for f in args.fmt.split(",") if f.strip()]
    exported = export_report(md_path.read_text(encoding="utf-8"), md_path.stem, Path(args.out), fmts)
    for fmt, p in exported.items():
        print(f"已导出 {fmt.upper()}：{p}")
    return 0


if __name__ == "__main__":
    sys.exit(main())